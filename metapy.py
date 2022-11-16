# Imports
import os
from PyPDF2 import PdfFileReader,PdfReader,PdfWriter #XXX Check pdfFileReader
import docx

##################################################################
#TODO                   Header
##################################################################
author = "Nacho Mestre"
version = 1.0
header = f"""
 __       __              __                _______
/  \     /  |            /  |              /       \.
$$  \   /$$ |  ______   _$$ |_     ______  $$$$$$$  | __    __
$$$  \ /$$$ | /      \ / $$   |   /      \ $$ |__$$ |/  |  /  |
$$$$  /$$$$ |/$$$$$$  |$$$$$$/    $$$$$$  |$$    $$/ $$ |  $$ |
$$ $$ $$/$$ |$$    $$ |  $$ | __  /    $$ |$$$$$$$/  $$ |  $$ |
$$ |$$$/ $$ |$$$$$$$$/   $$ |/  |/$$$$$$$ |$$ |      $$ \__$$ |
$$ | $/  $$ |$$       |  $$  $$/ $$    $$ |$$ |      $$    $$ |
$$/      $$/  $$$$$$$/    $$$$/   $$$$$$$/ $$/        $$$$$$$ |
                                                     /  \__$$ |
                                                     $$    $$/
                                                      $$$$$$/
Author: {author}
Version: {version}
"""
#######################################################################################################################################################
#TODO                                                           Functions
#######################################################################################################################################################
def nav(msg=None):
    current_path = os.getcwd()
    loop = True
    print(f"""
    -------------
    |    NAV    |
    -------------
    """)
    if msg != None:
        print(f"{msg}")    
    while loop:
        print(f"Current path: {current_path}")
        options = []
        dir_map = []
        ind=0
        print('------------------------------')
        for x in os.listdir(current_path):
            options.append(x)
            if os.path.isdir(x):
                print(f"{ind} - [ FOLDER ] {x}")
                dir_map.append(True)
            else:
                print(f"{ind} - [ FILE ] {x}")
                dir_map.append(False)
            ind+=1
        print('------------------------------')
        print('".." to go back')
        r = input("Select option -> ")
        if (r==".."):
            print("A MEDIAS")
        elif (r!='' and dir_map[int(r)]==False):
            return current_path+"/"+options[int(r)]
        elif (r!='' and dir_map[int(r)]==True):
            current_path += '/'+options[int(r)]
        else:
            return None

#######################################################################################################################################################
#TODO                                                           Extensions
#######################################################################################################################################################

#PDF -> https://pypdf2.readthedocs.io/en/latest/user/metadata.html
def pdf_read_meta(filePath):
    pdfFile = PdfFileReader(open(filePath, 'rb'))
    meta = pdfFile.getDocumentInfo()
    print()
    for metaItem in meta:
        print(f'    {metaItem[1:]}: {meta[metaItem]}')
    print()
    return meta

def pdf_write_meta(filePath):
    rd = PdfReader(filePath)
    wr = PdfWriter()
    for page in rd.pages:
        wr.add_page(page)
    print("")
    fields = {
    "/Author":input(f"Author: "),
    "/Producer":input(f"Producer: "),
    "/NeedAppearances": input(f"NeedAppearances: "),
    }
    wr.add_metadata(fields)
    with open(filePath,'wb') as f:
        wr.write(f)

def pdf_clear_meta(filePath):
    rd = PdfReader(filePath)
    wr = PdfWriter()
    for page in rd.pages:
        wr.add_page(page)
    fields = {"/Author":"","/Producer":"","/NeedAppearances":""}
    wr.add_metadata(fields)
    with open(filePath,'wb') as f:
        wr.write(f)
    
#DOCX
# READ
def docx_read_meta(filePath):
    doc = docx.Document(filePath)
    metadata = {}
    properties = doc.core_properties
    print(f"""
    Version: {properties.version}
    Title: {properties.title}
    Autor: {properties.author}
    Created: {properties.created}
    Category: {properties.category}
    Identifier: {properties.identifier}
    Comments: {properties.comments}
    Content_status: {properties.content_status}
    Keyword: {properties.keywords}
    last_modified_by: {properties.last_modified_by}
    language: {properties.language}
    modified: {properties.modified}
    subject: {properties.subject}
    """)
    metadata["category"] = properties.category
    metadata["comments"] = properties.comments
    metadata["content_status"] = properties.content_status
    metadata["created"] = properties.created
    metadata["identifier"] = properties.identifier
    metadata["keywords"] = properties.keywords
    metadata["last_modified_by"] = properties.last_modified_by
    metadata["language"] = properties.language
    metadata["modified"] = properties.modified
    metadata["subject"] = properties.subject
    metadata["title"] = properties.title
    metadata["version"] = properties.version
    return metadata

#Write
def docx_write_meta(filePath):
    doc = docx.Document(filePath)
    properties = doc.core_properties
    print("[*] Left empty to stay with same value")

    c = input(f"Author: ({properties.author})")
    if c != '':
        doc.core_properties.author = c

    c = input(f"Category: ({properties.category})")
    if c != '':
        doc.core_properties.category = c

    c = input(f"Comments: ({properties.comments})")
    if c != '':
        doc.core_properties.comments = c

    c = input(f"Content_status: ({properties.content_status})")
    if c != '':
        doc.core_properties.content_status = c

    c = input(f"Created: ({properties.created})")
    if c != '':
        doc.core_properties.created = c

    c = input(f"Identifier: ({properties.identifier})")
    if c != '':
        doc.core_properties.identifier = c

    c = input(f"Keywords: ({properties.keywords})")
    if c != '':
        doc.core_properties.keywords = c

    c = input(f"Last modified by: ({properties.last_modified_by})")
    if c != '':
        doc.core_properties.last_modified_by = c

    c = input(f"Language: ({properties.language})")
    if c != '':
        doc.core_properties.language = c

    c = input(f"Modified: ({properties.modified})")
    if c != '':
        doc.core_properties.modified = c

    c = input(f"Subject: ({properties.subject})")
    if c != '':
        doc.core_properties.subject = c

    c = input(f"Title: ({properties.title})")
    if c != '':
        doc.core_properties.title = c

    c = input(f"Version: ({properties.version})")
    if c != '':
        doc.core_properties.version = c
    # Deleting the old One and replacing for new One.
    import os
    os.remove(filePath)
    doc.save(filePath)

#CLEAR
def docx_clear_meta(filePath):
    #XXX
    from datetime import datetime
    d = datetime.now()
    doc = docx.Document(filePath)
    c=''
    doc.core_properties.author = c
    doc.core_properties.category = c
    doc.core_properties.comments = c
    doc.core_properties.content_status = c
    doc.core_properties.created = d
    doc.core_properties.identifier = c
    doc.core_properties.keywords = c
    doc.core_properties.last_modified_by = c
    doc.core_properties.language = c
    doc.core_properties.modified = d
    doc.core_properties.subject = c
    doc.core_properties.title = c
    doc.core_properties.version = c
    from os import remove as rm
    rm(filePath)
    doc.save(filePath)

# Just To print empyt lines an clean the console
def br(n_lines):
    for l in range(0,n_lines):
        print()

def clear():
    os.system("cls")

def get_Extension(filePath):
    if (os.path.isfile(filePath)):
        filePath, file_ext =  os.path.splitext(filePath)
        return file_ext
    return None

#######################################################################################################################################################
#TODO                                                           Main ()
#######################################################################################################################################################
if __name__ == '__main__':
    clear()
    loop = True;
    while loop:
        print(header)
        print('------------------')
        print("1. View Metadata")
        print("2. Edit & Write Metadata")
        print("3. Clear Metadata")
        br(2)
        print("99.")
        print("00. To exit")
        br(2)
        option = input("SELECT ANY OPTION -> ")
        clear()
        #TODO PREVIEW
        if (option=='1'):
            item = nav()
            if item is None:
                continue
            extension = get_Extension(item)
            if (extension == '.pdf'):
                pdf_read_meta(item)
            elif (extension == '.docx'):
                docx_read_meta(item)
            elif (extension == None):
                print("File doesnt exists")
            else:
                print("Invalid format ....")
            input('Press some key to continue.....')

        #TODO EDIT
        elif (option=='2'):
            item = nav("Select file to edit metainfo")
            if item is None:
                continue
            extension = get_Extension(item)
            if (extension == '.docx'):
                docx_write_meta(item)
                docx_read_meta(item)
            elif (extension == '.pdf'):
                pdf_write_meta(item)
            else:
                input("Extension not accepted...")
        #TODO CLEAR
        elif (option=='3'):
            item = nav("[WARNING] File path to remove metainfo ")#input("[WARNING] File path to remove metainfo -> ")
            if item is None:
                continue
            extension = get_Extension(item)
            if (extension == '.docx'):
                docx_clear_meta(item)

        elif (option == '00'):
            exit(0)

        else:
            input("Not valid option....")
        
        clear()
